import os
import openpyxl
import xlrd
from xlrd import xldate_as_tuple
import datetime
from docx import Document

class FileReader:
    """文件读取器，负责读取不同格式的文件"""
    
    @staticmethod
    def get_sheets(file_path):
        """获取文件的工作表列表，支持xlsx、xls、docx和doc格式
        
        Args:
            file_path (str): 文件路径
        
        Returns:
            list: 工作表名称列表
        
        Raises:
            Exception: 读取工作表失败时抛出异常
        """
        ext = os.path.splitext(file_path)[1].lower()
        sheets = []
        
        try:
            if ext == '.xlsx':
                # 使用openpyxl处理xlsx文件
                workbook = openpyxl.load_workbook(file_path)
                sheets = workbook.sheetnames
                workbook.close()
            elif ext == '.xls':
                # 使用xlrd处理xls文件
                workbook = xlrd.open_workbook(file_path)
                sheets = workbook.sheet_names()
            elif ext in ['.docx', '.doc']:
                # Word文件只有一个"工作表"（整个文档）
                sheets = ["文档内容"]
            else:
                raise Exception(f"不支持的文件格式: {ext}")
            
            return sheets
        except Exception as e:
            raise Exception(f"读取工作表失败: {str(e)}")
    
    @staticmethod
    def read_file_rows(file_path, sheet_name):
        """读取文件的行数据，支持xlsx、xls、docx和doc格式
        
        Args:
            file_path (str): 文件路径
            sheet_name (str): 工作表名称
        
        Returns:
            list: 行数据列表
        
        Raises:
            Exception: 读取文件失败时抛出异常
        """
        ext = os.path.splitext(file_path)[1].lower()
        rows = []
        
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                raise Exception(f"文件不存在: {file_path}")
            
            if ext == '.xlsx':
                try:
                    # 使用openpyxl处理xlsx文件
                    workbook = openpyxl.load_workbook(file_path)
                    sheet = workbook[sheet_name]
                    for row in sheet.iter_rows(min_row=1, values_only=True):
                        rows.append(row)
                    workbook.close()
                except Exception as e:
                    raise Exception(f"读取.xlsx文件失败: {str(e)}. 请检查文件是否损坏或格式不正确。")
            elif ext == '.xls':
                try:
                    # 使用xlrd处理xls文件
                    workbook = xlrd.open_workbook(file_path)
                    sheet = workbook.sheet_by_name(sheet_name)
                    for i in range(sheet.nrows):
                        row = []
                        for j in range(sheet.ncols):
                            cell_value = sheet.cell_value(i, j)
                            cell_type = sheet.cell_type(i, j)
                            
                            # 处理日期类型
                            if cell_type == xlrd.XL_CELL_DATE:
                                date_tuple = xldate_as_tuple(cell_value, workbook.datemode)
                                cell_value = datetime.datetime(*date_tuple).strftime('%Y-%m-%d')
                            # 处理数字类型，转换为字符串
                            elif cell_type == xlrd.XL_CELL_NUMBER:
                                # 如果是整数，转换为整数字符串，否则保留原格式
                                if cell_value == int(cell_value):
                                    cell_value = str(int(cell_value))
                                else:
                                    cell_value = str(cell_value)
                            # 处理空值
                            elif cell_type == xlrd.XL_CELL_EMPTY:
                                cell_value = ""
                            
                            row.append(cell_value)
                        rows.append(tuple(row))
                except Exception as e:
                    raise Exception(f"读取.xls文件失败: {str(e)}. 请检查文件是否损坏或格式不正确。")
            elif ext == '.docx':
                try:
                    # 使用python-docx处理docx文件
                    doc = Document(file_path)
                    # 读取文档内容，转换为类似Excel行的格式
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            # 将每个段落作为一行，只有一列
                            rows.append((text,))
                    # 读取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            rows.append(tuple(cells))
                except Exception as e:
                    raise Exception(f"读取.docx文件失败: {str(e)}. 请检查文件是否损坏或格式不正确。")
            elif ext == '.doc':
                # 使用win32com.client处理doc文件（Windows系统）
                try:
                    import win32com.client
                    # 初始化Word应用程序
                    word = win32com.client.Dispatch('Word.Application')
                    word.Visible = False
                    # 打开doc文件
                    doc = word.Documents.Open(file_path)
                    # 读取文档内容
                    text = doc.Content.Text
                    # 关闭文档
                    doc.Close()
                    # 退出Word应用程序
                    word.Quit()
                    # 将文本按换行符分割成段落
                    paragraphs = text.split('\n')
                    for para in paragraphs:
                        text = para.strip()
                        if text:
                            # 将每个段落作为一行，只有一列
                            rows.append((text,))
                except ImportError:
                    raise Exception("处理.doc文件需要安装pywin32库，请使用pip install pywin32命令安装")
                except Exception as e:
                    raise Exception(f"读取.doc文件失败: {str(e)}. 请检查文件是否损坏或格式不正确。")
            else:
                raise Exception(f"不支持的文件格式: {ext}. 请选择.xlsx、.xls、.docx或.doc格式的文件。")
            
            return rows
        except Exception as e:
            raise Exception(f"读取文件失败: {str(e)}")
