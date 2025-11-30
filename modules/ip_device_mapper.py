import os
import re
from docx import Document
from constants import IP_PATTERN, IP_DEVICE_PATTERNS

class IPDeviceMapper:
    """IP设备映射器，负责从Word文件中读取IP设备映射"""
    
    def __init__(self, logger=None):
        """初始化IP设备映射器
        
        Args:
            logger (callable, optional): 日志记录函数，默认None
        """
        self.logger = logger
    
    def log(self, message):
        """记录日志
        
        Args:
            message (str): 日志消息
        """
        if self.logger:
            self.logger(message)
    
    def read_ip_device_mapping(self, doc_path):
        """读取Word文件中的IP设备名称映射表，支持docx和doc格式
        
        Args:
            doc_path (str): 文档路径
        
        Returns:
            dict: IP设备映射字典
        
        Raises:
            Exception: 读取IP设备映射表失败时抛出异常
        """
        try:
            self.log(f"开始读取IP设备映射表: {doc_path}")
            
            # 检查文件是否存在
            if not os.path.exists(doc_path):
                raise Exception(f"文件不存在: {doc_path}")
            
            # 初始化映射字典
            ip_device_map = {}
            
            # 获取文件扩展名
            ext = os.path.splitext(doc_path)[1].lower()
            
            if ext == '.docx':
                ip_device_map = self._read_ip_device_mapping_from_docx(doc_path)
            elif ext == '.doc':
                ip_device_map = self._read_ip_device_mapping_from_doc(doc_path)
            else:
                raise Exception(f"不支持的文件格式: {ext}. 请选择.docx或.doc格式的文件。")
            
            self.log(f"\n读取完成，共找到 {len(ip_device_map)} 个IP设备映射")
            return ip_device_map
            
        except Exception as e:
            self.log(f"读取IP设备映射表失败: {str(e)}")
            raise Exception(f"读取IP设备映射表失败: {str(e)}")
    
    def _read_ip_device_mapping_from_docx(self, docx_path):
        """从docx文件读取IP设备映射
        
        Args:
            docx_path (str): 文档路径
        
        Returns:
            dict: IP设备映射字典
        
        Raises:
            Exception: 读取.docx文件失败时抛出异常
        """
        try:
            doc = Document(docx_path)
            ip_device_map = {}
            
            # 遍历所有段落
            self.log("=== 遍历段落 ===")
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    self._match_ip_device_patterns(text, ip_device_map)
            
            # 遍历所有表格
            self.log("\n=== 遍历表格 ===")
            for table_idx, table in enumerate(doc.tables):
                self.log(f"表格 {table_idx+1}，共 {len(table.rows)} 行，{len(table.columns)} 列")
                
                # 遍历表格的每一行
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    self._extract_ip_device_from_table_row(cells, ip_device_map)
            
            return ip_device_map
        except Exception as e:
            raise Exception(f"读取.docx文件失败: {str(e)}. 请检查文件是否损坏或格式不正确。")
    
    def _read_ip_device_mapping_from_doc(self, doc_path):
        """从doc文件读取IP设备映射
        
        Args:
            doc_path (str): 文档路径
        
        Returns:
            dict: IP设备映射字典
        
        Raises:
            Exception: 读取.doc文件失败时抛出异常
        """
        try:
            import win32com.client
            # 初始化Word应用程序
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            
            # 打开doc文件
            doc = word.Documents.Open(doc_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            
            # 将文本按换行符分割成段落
            paragraphs = text.split('\n')
            ip_device_map = {}
            
            # 遍历所有段落
            self.log("=== 遍历段落 ===")
            for para in paragraphs:
                text = para.strip()
                if text:
                    self._match_ip_device_patterns(text, ip_device_map)
            
            return ip_device_map
        except ImportError:
            raise Exception("处理.doc文件需要安装pywin32库，请使用pip install pywin32命令安装")
        except Exception as e:
            raise Exception(f"读取.doc文件失败: {str(e)}. 请检查文件是否损坏或格式不正确。")
    
    def _match_ip_device_patterns(self, text, ip_device_map):
        """匹配IP设备名称模式
        
        Args:
            text (str): 文本
            ip_device_map (dict): IP设备映射字典
        """
        for pattern in IP_DEVICE_PATTERNS:
            match = re.match(pattern, text)
            if match:
                ip = match.group(1)
                device = match.group(2).strip()
                ip_device_map[ip] = device
                self.log(f"  匹配到: IP={ip}, 设备名称={device}")
                break
    
    def _extract_ip_device_from_table_row(self, cells, ip_device_map):
        """从表格行提取IP设备映射
        
        Args:
            cells (list): 表格行单元格列表
            ip_device_map (dict): IP设备映射字典
        """
        # 检查是否至少有2个单元格
        if len(cells) >= 2:
            # 首先检查从右到左：设备名称 -> IP
            for i in range(1, len(cells)):
                # 检查当前单元格是否是IP地址
                if re.match(r'^' + IP_PATTERN + r'$', cells[i]):
                    # 当前单元格是IP地址，前一个单元格是设备名称
                    device_candidate = cells[i-1]
                    if device_candidate and device_candidate != cells[i]:
                        ip = cells[i]
                        device = device_candidate
                        ip_device_map[ip] = device
                        self.log(f"  匹配到: IP={ip}, 设备名称={device}")
                        return
            
            # 如果从右到左没有匹配到，再尝试从左到右：IP -> 设备名称
            for i in range(len(cells) - 1):
                # 检查当前单元格是否是IP地址
                if re.match(r'^' + IP_PATTERN + r'$', cells[i]):
                    # 当前单元格是IP地址，下一个单元格是设备名称
                    device_candidate = cells[i+1]
                    if device_candidate and device_candidate != cells[i]:
                        ip = cells[i]
                        device = device_candidate
                        ip_device_map[ip] = device
                        self.log(f"  匹配到: IP={ip}, 设备名称={device}")
                        return
